from __future__ import annotations

import json
import mimetypes
import os
import re
import shutil
import subprocess
import tempfile
from pathlib import Path
from urllib.parse import quote
from uuid import uuid4

from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, RedirectResponse
from pydantic import BaseModel, Field

from .aura_agent_core import AuraAgent
from .aura_chat_store import AuraChatStore, sha256_file
from .branding import ENDORSEMENT, PRODUCT_FULL_NAME, TAGLINE
from .plans import AURA_SPEECH
from .project import ProjectWorkspace
from .speech import AuraSpeechService
from .tenant_storage import list_project_dirs
from .web_access import AuraWebGateway

router = APIRouter(tags=["Aura Intelligence"])
store = AuraChatStore()
agent = AuraAgent(store=store)


class MessageRequest(BaseModel):
    message: str = Field(min_length=1, max_length=50000)
    attachment_ids: list[str] = Field(default_factory=list, max_length=12)


class ThreadPatch(BaseModel):
    title: str | None = Field(default=None, max_length=180)
    project_name: str | None = Field(default=None, max_length=120)
    clear_project: bool = False
    web_enabled: bool | None = None
    tools_enabled: bool | None = None
    voice_reply: bool | None = None


class EditMessageRequest(BaseModel):
    message: str = Field(min_length=1, max_length=50000)


class MemoryRequest(BaseModel):
    label: str = Field(default="Memory", max_length=120)
    content: str = Field(min_length=1, max_length=5000)


def _member(request: Request):
    member = getattr(request.state, "member", None)
    if member is None:
        raise HTTPException(401, "Sign in required")
    return member


def _public_attachment(item: dict) -> dict:
    return {
        "id": item.get("id"),
        "name": item.get("name"),
        "mime_type": item.get("mime_type"),
        "kind": item.get("kind"),
        "bytes": item.get("bytes"),
        "sha256": item.get("sha256"),
        "metadata": item.get("metadata") or {},
        "created_at": item.get("created_at"),
    }


def _thread_payload(user_id: str, thread_id: str) -> dict:
    thread = store.thread(user_id, thread_id)
    if not thread:
        raise HTTPException(404, "Aura conversation not found")
    messages = store.messages(user_id, thread_id, limit=400)
    rows = []
    for message in messages:
        rows.append(
            {
                **message,
                "attachments": [
                    _public_attachment(item)
                    for item in store.message_attachments(user_id, thread_id, message["id"])
                ],
            }
        )
    return {
        "thread": thread,
        "messages": rows,
        "summary": store.summary(user_id, thread_id),
        "tool_runs": store.tool_runs(user_id, thread_id, limit=80),
    }


@router.get("/aura-intelligence/api/status")
def aura_status(request: Request):
    member = _member(request)
    model_status = agent.diagnostics()
    try:
        web = AuraWebGateway().diagnostics()
        web_status = {
            "enabled": bool(web.get("enabled")),
            "search_configured": bool(web.get("self_hosted_search_configured")),
            "direct_https_fetch": bool(web.get("direct_https_fetch")),
            "private_network_fetch_blocked": bool(web.get("private_network_fetch_blocked")),
        }
    except Exception:
        web_status = {"enabled": False, "search_configured": False}
    speech = AuraSpeechService().diagnostics()
    return {
        "brand": PRODUCT_FULL_NAME,
        "member_plan": member.plan.id,
        "model": model_status,
        "web": web_status,
        "speech": {
            "available_to_plan": member.plan.has(AURA_SPEECH),
            "stt_configured": bool(speech.get("stt_command_configured") or (speech.get("whisper_cli") and speech.get("whisper_model_configured"))),
            "tts_configured": bool(speech.get("tts_command_configured") or speech.get("tts_url_configured") or speech.get("piper_model_configured")),
        },
        "features": {
            "persistent_threads": True,
            "conversation_search": True,
            "edit_and_regenerate": True,
            "branch_conversation": True,
            "project_pinning": True,
            "explicit_memory": True,
            "attachments": True,
            "voice_input": True,
            "voice_output": True,
            "tool_execution": True,
            "web_research": True,
            "automatic_long_chat_summary": True,
        },
    }


@router.get("/aura-intelligence/api/projects")
def aura_projects(request: Request):
    _member(request)
    rows = []
    for path in list_project_dirs():
        row = {
            "project_name": path.name,
            "title": path.name,
            "has_song_dna": (path / "song_dna.json").is_file(),
            "has_daw": (path / "aura_session.json").is_file(),
            "has_creative_manifest": (path / "creative_manifest.json").is_file(),
        }
        try:
            manifest = ProjectWorkspace(path).load_manifest()
            row["title"] = manifest.title
            row["mode"] = manifest.mode
        except Exception:
            pass
        rows.append(row)
    return rows


@router.get("/aura-intelligence/api/threads")
def list_threads(request: Request, q: str = "", limit: int = 100):
    member = _member(request)
    return store.list_threads(member.user_id, limit=limit, query=q)


@router.post("/aura-intelligence/api/threads")
def create_thread(request: Request):
    member = _member(request)
    return store.create_thread(member.user_id)


@router.get("/aura-intelligence/api/threads/{thread_id}")
def get_thread(thread_id: str, request: Request):
    member = _member(request)
    return _thread_payload(member.user_id, thread_id)


@router.patch("/aura-intelligence/api/threads/{thread_id}")
def patch_thread(thread_id: str, body: ThreadPatch, request: Request):
    member = _member(request)
    try:
        if body.title is not None:
            store.rename_thread(member.user_id, thread_id, body.title)
        current = store.thread(member.user_id, thread_id)
        if not current:
            raise KeyError(thread_id)
        project_value = current.get("project_name")
        if body.clear_project:
            project_value = ""
        elif body.project_name is not None:
            requested = body.project_name.strip()
            if requested:
                owned = {path.name for path in list_project_dirs()}
                if requested not in owned:
                    raise ValueError("Pinned project is not available to this member")
            project_value = requested
        store.set_context(
            member.user_id,
            thread_id,
            project_name=project_value,
            web_enabled=body.web_enabled,
            tools_enabled=body.tools_enabled,
            voice_reply=body.voice_reply,
        )
        return store.thread(member.user_id, thread_id)
    except KeyError as exc:
        raise HTTPException(404, "Aura conversation not found") from exc
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc


@router.delete("/aura-intelligence/api/threads/{thread_id}")
def delete_thread(thread_id: str, request: Request):
    member = _member(request)
    try:
        store.delete_thread(member.user_id, thread_id)
    except KeyError as exc:
        raise HTTPException(404, "Aura conversation not found") from exc
    return {"deleted": True, "thread_id": thread_id}


@router.post("/aura-intelligence/api/threads/{thread_id}/messages")
def send_message(thread_id: str, body: MessageRequest, request: Request):
    member = _member(request)
    try:
        return agent.respond(
            member=member,
            thread_id=thread_id,
            text=body.message,
            attachment_ids=body.attachment_ids,
        )
    except KeyError as exc:
        raise HTTPException(404, "Aura conversation not found") from exc
    except PermissionError as exc:
        raise HTTPException(403, str(exc)) from exc
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@router.post("/aura-intelligence/api/threads/{thread_id}/regenerate")
def regenerate_message(thread_id: str, request: Request):
    member = _member(request)
    try:
        return agent.regenerate(member=member, thread_id=thread_id)
    except KeyError as exc:
        raise HTTPException(404, "Aura conversation not found") from exc
    except ValueError as exc:
        raise HTTPException(409, str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@router.patch("/aura-intelligence/api/threads/{thread_id}/messages/{message_id}")
def edit_message(thread_id: str, message_id: str, body: EditMessageRequest, request: Request):
    member = _member(request)
    try:
        store.edit_user_message(member.user_id, thread_id, message_id, body.message)
        return agent.regenerate(member=member, thread_id=thread_id)
    except KeyError as exc:
        raise HTTPException(404, "User message not found") from exc
    except ValueError as exc:
        raise HTTPException(400, str(exc)) from exc
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc


@router.post("/aura-intelligence/api/threads/{thread_id}/branch/{message_id}")
def branch_thread(thread_id: str, message_id: str, request: Request):
    member = _member(request)
    try:
        return store.fork_thread(member.user_id, thread_id, message_id)
    except KeyError as exc:
        raise HTTPException(404, "Conversation/message not found") from exc


_TEXT_EXTS = {
    ".txt", ".md", ".json", ".yaml", ".yml", ".csv", ".tsv", ".py", ".js", ".ts", ".tsx", ".jsx",
    ".html", ".css", ".xml", ".log", ".ini", ".toml", ".sql", ".sh", ".ps1", ".java", ".c", ".cpp",
    ".h", ".hpp", ".go", ".rs", ".rb", ".php",
}
_IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp"}
_AUDIO_EXTS = {".wav", ".flac", ".mp3", ".m4a", ".ogg", ".aac", ".webm"}
_VIDEO_EXTS = {".mp4", ".mov", ".mkv", ".webm", ".m4v"}
_DOCUMENT_EXTS = {".pdf", ".docx", ".xlsx"}
_ALLOWED_ATTACHMENTS = _TEXT_EXTS | _IMAGE_EXTS | _AUDIO_EXTS | _VIDEO_EXTS | _DOCUMENT_EXTS


def _attachment_root(user_id: str, thread_id: str) -> Path:
    root = Path(os.getenv("AURA_CHAT_ATTACHMENT_DIR", "data/aura/attachments")).resolve()
    target = (root / user_id / thread_id).resolve()
    if root not in target.parents:
        raise ValueError("Invalid Aura attachment path")
    target.mkdir(parents=True, exist_ok=True)
    return target


def _attachment_kind(suffix: str) -> str:
    if suffix in _TEXT_EXTS:
        return "text"
    if suffix in _IMAGE_EXTS:
        return "image"
    if suffix in _AUDIO_EXTS:
        return "audio"
    if suffix in _VIDEO_EXTS:
        return "video"
    return "document"


def _extract_attachment(path: Path, kind: str) -> tuple[str | None, dict]:
    suffix = path.suffix.lower()
    metadata: dict = {}
    if kind == "text":
        text = path.read_text(encoding="utf-8", errors="replace")
        return text[:80000], {"characters": len(text), "truncated": len(text) > 80000}
    if suffix == ".pdf":
        try:
            from pypdf import PdfReader
            reader = PdfReader(str(path))
            parts = []
            for page in reader.pages[:80]:
                parts.append(page.extract_text() or "")
            text = "\n\n".join(parts).strip()
            return text[:80000], {"pages": len(reader.pages), "characters": len(text), "truncated": len(text) > 80000}
        except Exception as exc:
            return None, {"extraction_unavailable": f"{type(exc).__name__}: {exc}"}
    if suffix == ".docx":
        try:
            from docx import Document
            doc = Document(str(path))
            text = "\n".join(p.text for p in doc.paragraphs if p.text).strip()
            return text[:80000], {"paragraphs": len(doc.paragraphs), "characters": len(text), "truncated": len(text) > 80000}
        except Exception as exc:
            return None, {"extraction_unavailable": f"{type(exc).__name__}: {exc}"}
    if suffix == ".xlsx":
        try:
            from openpyxl import load_workbook
            wb = load_workbook(path, read_only=True, data_only=True)
            rows = []
            for ws in wb.worksheets[:12]:
                rows.append(f"[Sheet: {ws.title}]")
                for row in ws.iter_rows(max_row=200, values_only=True):
                    values = ["" if value is None else str(value) for value in row[:30]]
                    if any(values):
                        rows.append("\t".join(values))
            text = "\n".join(rows)
            return text[:80000], {"sheets": wb.sheetnames, "characters": len(text), "truncated": len(text) > 80000}
        except Exception as exc:
            return None, {"extraction_unavailable": f"{type(exc).__name__}: {exc}"}
    if kind == "audio":
        try:
            import soundfile as sf
            info = sf.info(path)
            metadata.update({
                "duration_seconds": float(info.frames / info.samplerate),
                "sample_rate": int(info.samplerate),
                "channels": int(info.channels),
            })
        except Exception:
            pass
        return None, metadata
    if kind == "image":
        try:
            from PIL import Image
            with Image.open(path) as image:
                metadata.update({"width": image.width, "height": image.height, "format": image.format})
        except Exception:
            pass
        return None, metadata
    if kind == "video" and shutil.which("ffprobe"):
        try:
            proc = subprocess.run(
                ["ffprobe", "-v", "quiet", "-print_format", "json", "-show_format", "-show_streams", str(path)],
                capture_output=True, text=True, check=True, timeout=30,
            )
            data = json.loads(proc.stdout)
            video = next((x for x in data.get("streams", []) if x.get("codec_type") == "video"), {})
            metadata.update({
                "duration_seconds": float((data.get("format") or {}).get("duration") or 0),
                "width": video.get("width"), "height": video.get("height"), "codec": video.get("codec_name"),
                "has_audio": any(x.get("codec_type") == "audio" for x in data.get("streams", [])),
            })
        except Exception:
            pass
    return None, metadata


@router.post("/aura-intelligence/api/threads/{thread_id}/attachments")
async def upload_attachment(thread_id: str, request: Request, file: UploadFile = File(...)):
    member = _member(request)
    if not store.thread(member.user_id, thread_id):
        raise HTTPException(404, "Aura conversation not found")
    original = Path(file.filename or "attachment").name[:240]
    suffix = Path(original).suffix.lower()
    if suffix not in _ALLOWED_ATTACHMENTS:
        raise HTTPException(415, "This attachment type is not supported by Aura chat yet")
    try:
        maximum_mb = max(1, min(500, int(os.getenv("AURA_CHAT_ATTACHMENT_MAX_MB", "50"))))
    except Exception:
        maximum_mb = 50
    maximum = maximum_mb * 1024 * 1024
    attachment_id = uuid4().hex
    safe_name = re.sub(r"[^A-Za-z0-9._-]+", "_", original).strip("._") or f"attachment{suffix}"
    target = _attachment_root(member.user_id, thread_id) / f"{attachment_id}_{safe_name}"
    total = 0
    try:
        with target.open("wb") as handle:
            while True:
                chunk = await file.read(1024 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > maximum:
                    raise HTTPException(413, f"Aura chat attachment exceeds the configured {maximum_mb} MB limit")
                handle.write(chunk)
    except Exception:
        target.unlink(missing_ok=True)
        raise
    finally:
        await file.close()
    kind = _attachment_kind(suffix)
    extracted, metadata = _extract_attachment(target, kind)
    item = store.add_attachment(
        member.user_id,
        thread_id,
        name=original,
        stored_path=str(target),
        mime_type=mimetypes.guess_type(original)[0] or file.content_type,
        kind=kind,
        bytes_count=total,
        sha256=sha256_file(target),
        extracted_text=extracted,
        metadata=metadata,
    )
    return _public_attachment(item)


@router.get("/aura-intelligence/api/memories")
def list_memories(request: Request):
    member = _member(request)
    return store.memories(member.user_id, enabled_only=False, limit=200)


@router.post("/aura-intelligence/api/memories")
def create_memory(body: MemoryRequest, request: Request):
    member = _member(request)
    return store.add_memory(member.user_id, body.label, body.content)


@router.delete("/aura-intelligence/api/memories/{memory_id}")
def delete_memory(memory_id: str, request: Request):
    member = _member(request)
    if not store.delete_memory(member.user_id, memory_id):
        raise HTTPException(404, "Aura memory not found")
    return {"deleted": True, "memory_id": memory_id}


@router.post("/aura-intelligence/api/threads/{thread_id}/voice-transcribe")
async def voice_transcribe(thread_id: str, request: Request, file: UploadFile = File(...)):
    member = _member(request)
    if not member.plan.has(AURA_SPEECH):
        raise HTTPException(403, "Aura voice input is not enabled for this membership")
    if not store.thread(member.user_id, thread_id):
        raise HTTPException(404, "Aura conversation not found")
    suffix = Path(file.filename or "voice.webm").suffix.lower() or ".webm"
    if suffix not in {".webm", ".wav", ".mp3", ".m4a", ".ogg", ".flac"}:
        raise HTTPException(415, "Unsupported voice-message audio format")
    maximum = 25 * 1024 * 1024
    total = 0
    temp_path = None
    try:
        with tempfile.NamedTemporaryFile(prefix="aura-voice-", suffix=suffix, delete=False) as temp:
            temp_path = Path(temp.name)
            while True:
                chunk = await file.read(512 * 1024)
                if not chunk:
                    break
                total += len(chunk)
                if total > maximum:
                    raise HTTPException(413, "Aura voice message exceeds 25 MB")
                temp.write(chunk)
        transcript = AuraSpeechService().transcribe(temp_path)
        return {"transcript": transcript, "bytes": total}
    except RuntimeError as exc:
        raise HTTPException(503, str(exc)) from exc
    finally:
        await file.close()
        if temp_path:
            temp_path.unlink(missing_ok=True)


@router.get("/aura-intelligence/api/threads/{thread_id}/messages/{message_id}/speech", include_in_schema=False)
def speak_message(thread_id: str, message_id: str, request: Request):
    member = _member(request)
    if not member.plan.has(AURA_SPEECH):
        raise HTTPException(403, "Aura speech output is not enabled for this membership")
    message = store.message(member.user_id, thread_id, message_id)
    if not message or message.get("role") != "assistant":
        raise HTTPException(404, "Aura response not found")
    root = Path(os.getenv("AURA_CHAT_SPEECH_DIR", "data/aura/speech")).resolve()
    target_dir = (root / member.user_id / thread_id).resolve()
    if root not in target_dir.parents:
        raise HTTPException(400, "Invalid speech output path")
    target_dir.mkdir(parents=True, exist_ok=True)
    target = target_dir / f"{message_id}.wav"
    if not target.is_file():
        try:
            AuraSpeechService().speak(message["content"][:12000], target)
        except RuntimeError as exc:
            raise HTTPException(503, str(exc)) from exc
    return FileResponse(target, media_type="audio/wav", headers={"Content-Disposition": "inline"})


CSS = r"""
:root{--bg:#03040a;--panel:#0b0f1d;--panel2:#11162a;--line:#ffffff18;--text:#f7f8ff;--muted:#aab3ca;--gold:#ebcb73;--violet:#9b73ff;--cyan:#58dcff;--green:#72e0a8;--red:#ff8fa5}*{box-sizing:border-box}html,body{margin:0;height:100%;background:radial-gradient(circle at 12% 0,#42186d55,transparent 27%),radial-gradient(circle at 95% 0,#143f7e55,transparent 25%),#03040a;color:var(--text);font-family:Inter,ui-sans-serif,system-ui,sans-serif}button,input,textarea,select{font:inherit}button{cursor:pointer}.app{height:100vh;display:grid;grid-template-columns:300px minmax(0,1fr)}.sidebar{border-right:1px solid var(--line);background:#050711e8;display:flex;flex-direction:column;min-height:0}.sideTop{padding:14px;border-bottom:1px solid var(--line)}.brand{font-weight:950;font-size:1.05rem}.brand small{display:block;color:var(--gold);font-size:.65rem;text-transform:uppercase;letter-spacing:.09em;margin-top:3px}.btn{border:1px solid var(--line);background:#ffffff08;color:#fff;border-radius:11px;padding:9px 11px;font-weight:800}.btn:hover{background:#ffffff10}.primary{border:0;background:linear-gradient(115deg,var(--gold),var(--violet));color:#160d1d}.new{width:100%;margin-top:11px}.search{width:100%;margin-top:9px;border:1px solid var(--line);border-radius:10px;background:#050812;color:#fff;padding:9px}.threads{overflow:auto;padding:8px;flex:1}.thread{border:1px solid transparent;border-radius:11px;padding:9px;margin:3px 0;cursor:pointer}.thread:hover,.thread.active{background:#ffffff09;border-color:var(--line)}.thread b{display:block;white-space:nowrap;overflow:hidden;text-overflow:ellipsis;font-size:.86rem}.thread small{color:var(--muted);font-size:.7rem}.sideFoot{padding:11px;border-top:1px solid var(--line);display:grid;gap:7px}.main{min-width:0;display:flex;flex-direction:column;height:100vh}.top{min-height:70px;border-bottom:1px solid var(--line);background:#050711dd;backdrop-filter:blur(16px);display:flex;align-items:center;gap:9px;padding:10px 15px;flex-wrap:wrap}.top .title{font-weight:950;margin-right:auto}.top .title small{display:block;color:var(--muted);font-size:.7rem;font-weight:600}.select,.toggle{border:1px solid var(--line);background:#080b15;color:#fff;border-radius:10px;padding:8px}.toggle{display:flex;align-items:center;gap:6px;font-size:.78rem}.messages{flex:1;overflow:auto;padding:26px 16px 160px}.conversation{width:min(900px,100%);margin:auto}.msg{margin:18px 0}.msg.user{display:flex;justify-content:flex-end}.bubble{max-width:85%;border:1px solid var(--line);border-radius:18px;padding:13px 15px;line-height:1.55;white-space:normal;overflow-wrap:anywhere}.user .bubble{background:linear-gradient(145deg,#2a1d4e,#172440);border-color:#9b73ff44}.assistant .bubble{max-width:100%;background:transparent;border-color:transparent;padding-left:0}.meta{display:flex;gap:6px;align-items:center;margin-top:6px;opacity:.75;flex-wrap:wrap}.mini{border:0;background:transparent;color:var(--muted);padding:4px 6px;border-radius:7px;font-size:.72rem}.mini:hover{background:#ffffff0c;color:#fff}.attach{display:inline-flex;gap:6px;align-items:center;border:1px solid var(--line);border-radius:10px;padding:6px 8px;margin:5px 5px 0 0;color:var(--muted);font-size:.72rem}.composerWrap{position:fixed;left:300px;right:0;bottom:0;padding:14px 18px 18px;background:linear-gradient(transparent,#03040a 28%)}.composer{width:min(920px,100%);margin:auto;border:1px solid #ffffff24;border-radius:20px;background:#0a0e1bf5;box-shadow:0 16px 50px #0009;padding:9px}.pending{display:flex;gap:6px;flex-wrap:wrap;padding:2px 4px}.composeRow{display:flex;gap:7px;align-items:flex-end}.composer textarea{flex:1;resize:none;max-height:180px;min-height:48px;border:0;outline:0;background:transparent;color:#fff;padding:11px;line-height:1.45}.round{width:42px;height:42px;border-radius:50%;border:1px solid var(--line);background:#ffffff08;color:#fff;font-weight:900}.send{background:linear-gradient(115deg,var(--gold),var(--violet));color:#160d1d;border:0}.notice{position:fixed;right:18px;top:82px;z-index:30;max-width:420px;background:#111629;border:1px solid var(--line);border-radius:12px;padding:10px 12px;display:none}.notice.show{display:block}.drawer{position:fixed;right:0;top:0;bottom:0;width:min(430px,100%);background:#090d19f8;border-left:1px solid var(--line);z-index:40;transform:translateX(105%);transition:.2s;padding:16px;overflow:auto}.drawer.open{transform:translateX(0)}.mem{border:1px solid var(--line);border-radius:12px;padding:10px;margin:8px 0;background:#ffffff05}.muted{color:var(--muted)}pre{white-space:pre-wrap;background:#070a12;border:1px solid var(--line);border-radius:10px;padding:11px;overflow:auto}code{font-family:ui-monospace,SFMono-Regular,Menlo,monospace}.empty{text-align:center;color:var(--muted);padding:80px 20px}.tools{font-size:.72rem;color:var(--green);margin-top:5px}.mobileMenu{display:none}@media(max-width:800px){.app{grid-template-columns:1fr}.sidebar{position:fixed;left:0;top:0;bottom:0;width:min(310px,90vw);z-index:50;transform:translateX(-105%);transition:.2s}.sidebar.open{transform:translateX(0)}.composerWrap{left:0}.mobileMenu{display:inline-block}.top{padding:9px}.bubble{max-width:96%}}
"""


SCRIPT = r"""
const A='/aura-intelligence/api';let current=null,threads=[],projects=[],pending=[],recording=false,recorder=null,chunks=[];const $=id=>document.getElementById(id);function esc(v){return String(v??'').replace(/[&<>'\"]/g,c=>({'&':'&amp;','<':'&lt;','>':'&gt;',"'":'&#39;','\"':'&quot;'}[c]))}function md(v){let x=esc(v);const blocks=[];x=x.replace(/```([\s\S]*?)```/g,(m,c)=>{blocks.push(`<pre><code>${c.trim()}</code></pre>`);return `@@CODE${blocks.length-1}@@`});x=x.replace(/`([^`]+)`/g,'<code>$1</code>').replace(/\*\*([^*]+)\*\*/g,'<b>$1</b>').replace(/\n/g,'<br>');blocks.forEach((b,i)=>x=x.replace(`@@CODE${i}@@`,b));return x}function note(m,b=false){const n=$('notice');n.textContent=m;n.style.borderColor=b?'#ff8fa555':'';n.className='notice show';clearTimeout(window._nt);window._nt=setTimeout(()=>n.className='notice',5000)}async function req(u,o={}){const r=await fetch(u,{credentials:'same-origin',headers:{'Content-Type':'application/json',...(o.headers||{})},...o});let b={};try{b=await r.json()}catch(e){}if(!r.ok)throw new Error(b.detail||`Request failed (${r.status})`);return b}function fmt(n){if(n<1024)return n+' B';if(n<1048576)return (n/1024).toFixed(1)+' KB';return (n/1048576).toFixed(1)+' MB'}
async function boot(){try{const [t,p,s]=await Promise.all([req(A+'/threads'),req(A+'/projects'),req(A+'/status')]);threads=t;projects=p;renderThreads();renderProjects();$('modelState').textContent=(s.model?.provider_mode||'Aura')+' · '+(s.web?.search_configured?'web ready':'web local');if(threads.length)await openThread(threads[0].id);else await newThread()}catch(e){note(e.message,true)}}function renderThreads(){const q=($('search').value||'').toLowerCase();$('threads').innerHTML=threads.filter(t=>!q||String(t.title).toLowerCase().includes(q)).map(t=>`<div class="thread ${current===t.id?'active':''}" onclick="openThread('${t.id}')"><b>${esc(t.title)}</b><small>${t.project_name?'📌 '+esc(t.project_name)+' · ':''}${t.message_count||0} messages</small></div>`).join('')||'<div class="empty">No conversations</div>'}function renderProjects(){const s=$('project');s.innerHTML='<option value="">No project pinned</option>'+projects.map(p=>`<option value="${esc(p.project_name)}">${esc(p.title)} · ${esc(p.project_name)}</option>`).join('')}
async function refreshThreads(){threads=await req(A+'/threads');renderThreads()}async function newThread(){const t=await req(A+'/threads',{method:'POST',body:'{}'});await refreshThreads();await openThread(t.id);$('sidebar').classList.remove('open')}async function openThread(id){current=id;const d=await req(`${A}/threads/${id}`);renderThread(d);renderThreads();$('sidebar').classList.remove('open')}function renderThread(d){$('chatTitle').textContent=d.thread.title;$('project').value=d.thread.project_name||'';$('web').checked=!!d.thread.web_enabled;$('tools').checked=!!d.thread.tools_enabled;const c=$('conversation');if(!d.messages.length){c.innerHTML='<div class="empty"><h2>Aura is ready.</h2><p>Ask anything, pin a project, attach a file, speak a request, research the web or tell Aura to perform a connected Studio action.</p></div>';return}const toolByMsg={};(d.tool_runs||[]).forEach(x=>{if(x.message_id)(toolByMsg[x.message_id]??=[]).push(x)});c.innerHTML=d.messages.map(m=>{const at=(m.attachments||[]).map(a=>`<span class="attach">📎 ${esc(a.name)} · ${fmt(a.bytes||0)}</span>`).join('');const tr=(toolByMsg[m.id]||[]).map(x=>`<span>${x.status==='completed'?'✓':'!'} ${esc(x.tool_name)}</span>`).join(' · ');const actions=m.role==='assistant'?`<button class="mini" onclick="copyMsg('${m.id}')">Copy</button><button class="mini" onclick="regen()">Regenerate</button><button class="mini" onclick="speak('${m.id}')">🔊 Speak</button><button class="mini" onclick="branch('${m.id}')">Branch</button>`:`<button class="mini" onclick="editMsg('${m.id}',${JSON.stringify(m.content)})">Edit</button><button class="mini" onclick="branch('${m.id}')">Branch</button>`;return `<div class="msg ${m.role}" data-id="${m.id}"><div class="bubble">${at}<div>${md(m.content)}</div>${tr?`<div class="tools">Aura tools: ${tr}</div>`:''}<div class="meta">${actions}</div></div></div>`}).join('');setTimeout(()=>{$('messages').scrollTop=$('messages').scrollHeight},30)}
async function sendText(text=null){if(!current)return;const box=$('composer');const message=(text??box.value).trim();if(!message)return;box.value='';const ids=pending.map(x=>x.id);pending=[];renderPending();box.disabled=true;try{await req(`${A}/threads/${current}/messages`,{method:'POST',body:JSON.stringify({message,attachment_ids:ids})});await refreshThreads();await openThread(current)}catch(e){note(e.message,true)}finally{box.disabled=false;box.focus()}}function keysend(e){if(e.key==='Enter'&&!e.shiftKey){e.preventDefault();sendText()}}async function regen(){try{await req(`${A}/threads/${current}/regenerate`,{method:'POST',body:'{}'});await openThread(current)}catch(e){note(e.message,true)}}async function editMsg(id,old){const next=prompt('Edit your message:',old);if(next===null||!next.trim())return;try{await req(`${A}/threads/${current}/messages/${id}`,{method:'PATCH',body:JSON.stringify({message:next.trim()})});await refreshThreads();await openThread(current)}catch(e){note(e.message,true)}}async function branch(id){try{const t=await req(`${A}/threads/${current}/branch/${id}`,{method:'POST',body:'{}'});await refreshThreads();await openThread(t.id);note('Conversation branched.')}catch(e){note(e.message,true)}}function copyMsg(id){const el=document.querySelector(`[data-id="${id}"] .bubble > div:nth-of-type(1)`);navigator.clipboard.writeText(el?.innerText||'').then(()=>note('Copied.'))}
async function setContext(){if(!current)return;try{await req(`${A}/threads/${current}`,{method:'PATCH',body:JSON.stringify({project_name:$('project').value,clear_project:!$('project').value,web_enabled:$('web').checked,tools_enabled:$('tools').checked})});await refreshThreads();note('Aura context updated.')}catch(e){note(e.message,true)}}async function renameThread(){if(!current)return;const t=threads.find(x=>x.id===current);const v=prompt('Conversation title:',t?.title||'');if(!v)return;try{await req(`${A}/threads/${current}`,{method:'PATCH',body:JSON.stringify({title:v})});await refreshThreads();$('chatTitle').textContent=v}catch(e){note(e.message,true)}}async function delThread(){if(!current||!confirm('Delete this Aura conversation?'))return;try{await req(`${A}/threads/${current}`,{method:'DELETE'});current=null;await refreshThreads();if(threads.length)openThread(threads[0].id);else newThread()}catch(e){note(e.message,true)}}
async function uploadFiles(files){if(!current)return;for(const file of files){const fd=new FormData();fd.append('file',file);try{const r=await fetch(`${A}/threads/${current}/attachments`,{method:'POST',credentials:'same-origin',body:fd});const b=await r.json();if(!r.ok)throw new Error(b.detail||'Upload failed');pending.push(b);renderPending()}catch(e){note(e.message,true)}}$('file').value=''}function renderPending(){$('pending').innerHTML=pending.map((a,i)=>`<span class="attach">📎 ${esc(a.name)} <button class="mini" onclick="pending.splice(${i},1);renderPending()">×</button></span>`).join('')}
async function mic(){if(recording){recorder.stop();return}if(!navigator.mediaDevices?.getUserMedia)return note('Microphone capture is not available in this browser.',true);try{const stream=await navigator.mediaDevices.getUserMedia({audio:true});chunks=[];recorder=new MediaRecorder(stream);recorder.ondataavailable=e=>{if(e.data.size)chunks.push(e.data)};recorder.onstop=async()=>{recording=false;$('mic').textContent='🎙';stream.getTracks().forEach(t=>t.stop());const blob=new Blob(chunks,{type:recorder.mimeType||'audio/webm'});const fd=new FormData();fd.append('file',blob,'aura-voice.webm');try{$('mic').textContent='…';const r=await fetch(`${A}/threads/${current}/voice-transcribe`,{method:'POST',credentials:'same-origin',body:fd});const b=await r.json();if(!r.ok)throw new Error(b.detail||'Voice transcription failed');$('composer').value=b.transcript;await sendText(b.transcript)}catch(e){note(e.message,true)}finally{$('mic').textContent='🎙'}};recorder.start();recording=true;$('mic').textContent='■'}catch(e){note('Microphone permission was not granted.',true)}}function speak(id){const a=new Audio(`${A}/threads/${current}/messages/${id}/speech`);a.play().catch(()=>note('Aura speech output is not configured on this host.',true))}
async function openMemory(){$('memory').classList.add('open');await loadMemory()}async function loadMemory(){try{const rows=await req(A+'/memories');$('memoryRows').innerHTML=rows.map(x=>`<div class="mem"><b>${esc(x.label)}</b><p>${esc(x.content)}</p><button class="mini" onclick="forget('${x.id}')">Delete memory</button></div>`).join('')||'<p class="muted">Aura has no saved memories. Ordinary conversations are not silently copied here.</p>'}catch(e){note(e.message,true)}}async function saveMemory(){const content=$('memText').value.trim();if(!content)return;try{await req(A+'/memories',{method:'POST',body:JSON.stringify({label:$('memLabel').value||'Memory',content})});$('memText').value='';await loadMemory()}catch(e){note(e.message,true)}}async function forget(id){if(!confirm('Delete this saved Aura memory?'))return;await req(`${A}/memories/${id}`,{method:'DELETE'});await loadMemory()}function side(){$('sidebar').classList.toggle('open')}$('search').addEventListener('input',async()=>{const q=$('search').value.trim();if(q.length>2){try{threads=await req(A+'/threads?q='+encodeURIComponent(q));renderThreads()}catch(e){}}else if(!q)refreshThreads()});boot();
"""


@router.get("/aura-intelligence", response_class=HTMLResponse, include_in_schema=False)
def aura_intelligence_page(request: Request):
    member = getattr(request.state, "member", None)
    if member is None:
        return RedirectResponse("/signin?next=/aura-intelligence", status_code=303)
    html = f"""<!doctype html><html lang='en'><head><meta charset='utf-8'><meta name='viewport' content='width=device-width,initial-scale=1'><meta name='robots' content='noindex,nofollow'><title>Aura — {PRODUCT_FULL_NAME}</title><style>{CSS}</style></head><body><div class='app'><aside class='sidebar' id='sidebar'><div class='sideTop'><div class='brand'>Aura<small>{PRODUCT_FULL_NAME}</small></div><button class='btn primary new' onclick='newThread()'>＋ New conversation</button><input class='search' id='search' placeholder='Search conversations'></div><div class='threads' id='threads'></div><div class='sideFoot'><button class='btn' onclick='openMemory()'>🧠 Aura Memory</button><a class='btn' href='/dashboard' style='text-align:center;text-decoration:none'>← Dashboard</a></div></aside><main class='main'><header class='top'><button class='btn mobileMenu' onclick='side()'>☰</button><div class='title'><span id='chatTitle'>Aura</span><small id='modelState'>Loading Aura Core…</small></div><select class='select' id='project' onchange='setContext()'></select><label class='toggle'><input type='checkbox' id='web' onchange='setContext()'> Web</label><label class='toggle'><input type='checkbox' id='tools' onchange='setContext()'> Tools</label><button class='btn' onclick='renameThread()'>Rename</button><button class='btn' onclick='delThread()'>Delete</button></header><section class='messages' id='messages'><div class='conversation' id='conversation'></div></section><div class='composerWrap'><div class='composer'><div class='pending' id='pending'></div><div class='composeRow'><input id='file' type='file' multiple hidden onchange='uploadFiles(this.files)'><button class='round' onclick="$('file').click()" title='Attach files'>＋</button><button class='round' id='mic' onclick='mic()' title='Speak to Aura'>🎙</button><textarea id='composer' rows='1' placeholder='Message Aura…' onkeydown='keysend(event)'></textarea><button class='round send' onclick='sendText()'>↑</button></div></div></div></main></div><div class='notice' id='notice'></div><aside class='drawer' id='memory'><div style='display:flex;justify-content:space-between;align-items:center'><div><b>Aura Memory</b><div class='muted' style='font-size:.75rem'>Only information you explicitly save belongs here.</div></div><button class='btn' onclick="$('memory').classList.remove('open')">×</button></div><input class='search' id='memLabel' placeholder='Memory label'><textarea class='search' id='memText' rows='4' placeholder='What should Aura remember?'></textarea><button class='btn primary' onclick='saveMemory()'>Save memory</button><div id='memoryRows'></div></aside><script>{SCRIPT}</script></body></html>"""
    return HTMLResponse(html)


__all__ = ["router"]
