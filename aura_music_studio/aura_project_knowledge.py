from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader

from . import aura_agent_core as core
from . import aura_agent_tools as tools
from .creative_project import CreativeProjectStore

_INSTALLED = False
_MAX_FILES = 120
_MAX_FILE_BYTES = 25 * 1024 * 1024
_MAX_TEXT_PER_FILE = 120_000
_TEXT_EXTS = {".txt", ".md", ".json", ".yaml", ".yml", ".csv", ".tsv", ".py", ".js", ".ts", ".html", ".css"}
_DOC_EXTS = _TEXT_EXTS | {".pdf", ".docx", ".xlsx", ".xlsm"}
_STOP = {"the", "and", "for", "with", "this", "that", "from", "into", "what", "does", "project", "files", "file", "about", "find", "search", "show", "tell"}


KNOWLEDGE_SPEC = tools.ToolSpec(
    name="search_project_knowledge",
    description=(
        "Search private documents and structured project metadata inside the pinned project. Supports promoted PDF/DOCX/XLSX/text plus Song DNA, Creative DNA and project manifests. Returns bounded excerpts and project-relative source refs, never raw server paths."
    ),
    arguments={
        "project_name": "Project name; omit when pinned.",
        "query": "What to find in the project knowledge.",
        "limit": "Optional 1-12 matches; default 8.",
    },
    write=False,
    web=False,
)


def _safe_file(project: Path, source_ref: str) -> Path | None:
    root = project.resolve()
    try:
        target = (root / source_ref).resolve()
    except Exception:
        return None
    if target == root or root not in target.parents or not target.is_file():
        return None
    if target.suffix.lower() not in _DOC_EXTS or target.stat().st_size > _MAX_FILE_BYTES:
        return None
    return target


def _candidates(project: Path) -> list[dict]:
    root = project.resolve()
    rows: list[dict] = []
    seen: set[str] = set()

    def add(path: Path, label: str | None = None):
        try:
            resolved = path.resolve()
            if root not in resolved.parents or not resolved.is_file() or resolved.suffix.lower() not in _DOC_EXTS:
                return
            if resolved.stat().st_size > _MAX_FILE_BYTES:
                return
            key = resolved.as_posix()
            if key in seen or len(rows) >= _MAX_FILES:
                return
            seen.add(key)
            rows.append({
                "path": resolved,
                "source_ref": resolved.relative_to(root).as_posix(),
                "label": label or resolved.name,
            })
        except Exception:
            return

    for name in ("project.yaml", "project.yml", "project.json", "song_dna.json", "creative_manifest.json", "aura_session.json"):
        add(project / name)

    creative = CreativeProjectStore(project)
    if creative.exists():
        try:
            for ref in creative.load().references:
                path = _safe_file(project, ref.source_ref)
                if path:
                    add(path, ref.label)
        except Exception:
            pass

    references = project / "input" / "references"
    if references.is_dir():
        for path in references.rglob("*"):
            add(path)
            if len(rows) >= _MAX_FILES:
                break
    return rows


def _text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")[:_MAX_TEXT_PER_FILE]


def _pdf_text(path: Path) -> str:
    reader = PdfReader(str(path))
    parts = []
    total = 0
    for page in reader.pages[:40]:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text:
            parts.append(text)
            total += len(text)
        if total >= _MAX_TEXT_PER_FILE:
            break
    return "\n".join(parts)[:_MAX_TEXT_PER_FILE]


def _docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables[:30]:
        for row in table.rows[:300]:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts)[:_MAX_TEXT_PER_FILE]


def _xlsx_text(path: Path) -> str:
    workbook = load_workbook(path, read_only=True, data_only=True)
    parts = []
    total = 0
    try:
        for sheet_name in workbook.sheetnames[:10]:
            parts.append(f"SHEET: {sheet_name}")
            ws = workbook[sheet_name]
            for index, row in enumerate(ws.iter_rows(values_only=True)):
                if index >= 500:
                    break
                line = " | ".join("" if value is None else str(value) for value in row[:100])
                parts.append(line)
                total += len(line)
                if total >= _MAX_TEXT_PER_FILE:
                    break
            if total >= _MAX_TEXT_PER_FILE:
                break
    finally:
        workbook.close()
    return "\n".join(parts)[:_MAX_TEXT_PER_FILE]


def _extract(path: Path) -> str:
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(path)
    if suffix == ".docx":
        return _docx_text(path)
    if suffix in {".xlsx", ".xlsm"}:
        return _xlsx_text(path)
    return _text_file(path)


def _terms(query: str) -> tuple[str, list[str]]:
    clean = " ".join((query or "").split()).strip()[:1000]
    words = [word for word in re.findall(r"[a-z0-9][a-z0-9_-]{2,}", clean.lower()) if word not in _STOP]
    unique = []
    for word in words:
        if word not in unique:
            unique.append(word)
    if not unique:
        raise ValueError("Project knowledge search needs a more specific query")
    return clean, unique[:20]


def _excerpt(text: str, terms: list[str], radius: int = 650) -> str:
    lower = text.lower()
    positions = [lower.find(term) for term in terms if lower.find(term) >= 0]
    start = max(0, (min(positions) if positions else 0) - radius)
    end = min(len(text), start + radius * 2)
    value = re.sub(r"\s+", " ", text[start:end]).strip()
    return value[:1400]


def search_project(project: Path, query: str, limit: int = 8) -> list[dict]:
    phrase, terms = _terms(query)
    phrase_lower = phrase.lower()
    matches = []
    for row in _candidates(project):
        try:
            text = _extract(row["path"])
        except Exception as exc:
            continue
        if not text:
            continue
        lower = text.lower()
        counts = {term: lower.count(term) for term in terms}
        matched_terms = [term for term, count in counts.items() if count]
        if not matched_terms:
            continue
        phrase_count = lower.count(phrase_lower) if len(phrase_lower) >= 5 else 0
        score = phrase_count * 20 + sum(min(count, 25) for count in counts.values()) + len(matched_terms) * 3
        matches.append({
            "source_ref": row["source_ref"],
            "title": row["label"],
            "kind": row["path"].suffix.lower().lstrip(".") or "text",
            "score": score,
            "matched_terms": matched_terms,
            "excerpt": _excerpt(text, matched_terms),
            "content_truncated_for_search": len(text) >= _MAX_TEXT_PER_FILE,
        })
    matches.sort(key=lambda item: (item["score"], len(item["matched_terms"])), reverse=True)
    return matches[: max(1, min(int(limit or 8), 12))]


def _looks_knowledge(text: str) -> bool:
    lower = (text or "").lower()
    return any(phrase in lower for phrase in (
        "search this project", "search the project", "find in this project", "find in the project",
        "search project files", "look through the project", "what do the project files say", "what does the project say",
        "search my project", "find this in my project",
    ))


def install_aura_project_knowledge() -> None:
    global _INSTALLED
    if _INSTALLED:
        return
    if KNOWLEDGE_SPEC.name not in {item.name for item in tools.TOOL_SPECS}:
        tools.TOOL_SPECS.append(KNOWLEDGE_SPEC)
        tools._SPEC_BY_NAME[KNOWLEDGE_SPEC.name] = KNOWLEDGE_SPEC
    original_execute = tools.AuraToolRegistry.execute
    original_direct = core._direct_tool_plan
    original_needs = core._needs_model_tool_router

    def execute(self, call: tools.ToolCall, *, latest_user_message: str):
        if call.name != "search_project_knowledge":
            return original_execute(self, call, latest_user_message=latest_user_message)
        if not self.tools_enabled:
            raise PermissionError("Aura tools are disabled for this conversation")
        args = dict(call.arguments or {})
        name = tools._project_name(args, self.pinned_project)
        project = tools._safe_project(name)
        query = str(args.get("query") or latest_user_message)
        results = search_project(project, query, int(args.get("limit") or 8))
        return {"project_name": name, "query": query, "matches": results, "raw_storage_paths_exposed": False, "local_private_search": True}

    def direct_tool_plan(text: str, pinned_project: str | None, web_enabled: bool):
        prior = original_direct(text, pinned_project, web_enabled)
        if prior is not None:
            return prior
        if pinned_project and _looks_knowledge(text):
            return tools.ToolPlan(calls=[tools.ToolCall(name="search_project_knowledge", arguments={"project_name": pinned_project, "query": text, "limit": 8})])
        return None

    def needs_model_tool_router(text: str, pinned_project: str | None, tools_enabled: bool, web_enabled: bool) -> bool:
        if tools_enabled and pinned_project and _looks_knowledge(text):
            return True
        return original_needs(text, pinned_project, tools_enabled, web_enabled)

    tools.AuraToolRegistry.execute = execute
    core._direct_tool_plan = direct_tool_plan
    core._needs_model_tool_router = needs_model_tool_router
    _INSTALLED = True


__all__ = ["install_aura_project_knowledge", "search_project"]
